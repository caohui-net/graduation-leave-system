#!/usr/bin/env python3
"""Convert Markdown user manual to DOCX format."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown(md_path):
    """Parse Markdown file and return structured content."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def add_heading(doc, text, level):
    """Add heading with appropriate style."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_paragraph(doc, text, style=None):
    """Add paragraph with optional style."""
    p = doc.add_paragraph(text, style=style)
    return p

def add_table_from_markdown(doc, table_text):
    """Parse markdown table and add to document."""
    lines = [l.strip() for l in table_text.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return

    # Parse header
    header = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    # Skip separator line
    # Parse rows
    rows_data = []
    for line in lines[2:]:
        row = [cell.strip() for cell in line.split('|') if cell.strip()]
        if row:
            rows_data.append(row)

    if not header or not rows_data:
        return

    # Create table
    table = doc.add_table(rows=len(rows_data)+1, cols=len(header))
    table.style = 'Light Grid Accent 1'

    # Add header
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(table.rows[row_idx+1].cells):
                table.rows[row_idx+1].cells[col_idx].text = cell_text

def convert_md_to_docx(md_path, docx_path):
    """Convert Markdown to DOCX."""
    content = parse_markdown(md_path)
    doc = Document()

    # Split content into lines
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    in_table = False
    table_content = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_content = []
            else:
                # End code block
                in_code_block = False
                p = doc.add_paragraph('\n'.join(code_block_content))
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # Table detection
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_content = [line]
            else:
                table_content.append(line)
            i += 1
            # Check if next line is still table
            if i < len(lines) and '|' in lines[i]:
                continue
            else:
                # End of table
                in_table = False
                add_table_from_markdown(doc, '\n'.join(table_content))
                table_content = []
                continue

        # Headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level <= 3:
                add_heading(doc, text, level)
        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)
        # Image placeholder
        elif line.strip().startswith('!['):
            match = re.match(r'!\[([^\]]+)\]\(([^)]+)\)', line.strip())
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                p = doc.add_paragraph()
                run = p.add_run(f'[截图占位符: {alt_text}]')
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.italic = True
                p = doc.add_paragraph()
                run = p.add_run(f'图片路径: {img_path}')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
        # Bold metadata
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
        # List items
        elif line.strip().startswith('-') or re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^[-\d]+\.?\s*', '', line.strip())
            doc.add_paragraph(text, style='List Bullet')
        # Regular paragraph
        elif line.strip():
            # Check for bold/italic inline
            if '**' in line or '*' in line:
                p = doc.add_paragraph()
                # Simple bold handling
                parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part.strip('*'))
                        run.font.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = p.add_run(part.strip('*'))
                        run.font.italic = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(line.strip())
        # Empty line - add spacing
        else:
            doc.add_paragraph()

        i += 1

    doc.save(docx_path)

if __name__ == '__main__':
    md_file = Path(__file__).parent.parent / 'docs' / '用户操作手册.md'
    docx_file = Path(__file__).parent.parent / 'docs' / '用户操作手册.docx'

    print(f"Converting {md_file} to {docx_file}...")
    convert_md_to_docx(md_file, docx_file)
    print(f"Conversion complete: {docx_file}")
